from crewai import Crew
from agents import researcher
from tasks import research_task
import tools
import streamlit as st
from docx import Document
from io import BytesIO


crew = Crew(
    agents=[researcher],
    tasks=[research_task],
    verbose=True
)

st.set_page_config(page_title="CrewAI Notes Generator", layout="centered")
st.title("CrewAI Agent for Notes Generator")

topic = st.text_input("Enter the topic to research : ").strip() 

if st.button("Generate Notes"):
    with st.spinner("Generating notes..."):
        result = crew.kickoff(inputs={"topic": topic.strip()})
    st.success("Research Completed ✅")
    st.subheader("📄 Result:")
    st.write(result)
    
    doc=Document()
    doc.add_heading(f"Notes on {topic}", level=1)
    doc.add_paragraph(str(result))
    
    buffer=BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.download_button(
        label="Download Notes",
        data=buffer,
        file_name=f"{topic}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
